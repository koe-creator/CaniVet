from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = BASE_DIR

PROJECT = {
    "nombre": "CaniVet - Sistema de Gestion Canina y Veterinaria",
    "empresa": "Proyecto academico / plataforma veterinaria CaniVet",
    "tipo": "Sistema web administrativo",
    "frontend": "React 19 + Vite",
    "backend": "Flask 3 + Python 3.12",
    "bd": "Supabase PostgreSQL",
    "auth": "Supabase Auth con control de roles",
    "correo": "Gmail SMTP desde backend Flask",
}


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_document_language(document, lang="es-DO"):
    styles = document.styles
    for style_name in ["Normal", "Title", "Subtitle", "Heading 1", "Heading 2", "Heading 3"]:
        if style_name in styles:
            style = styles[style_name]
            rpr = style.element.get_or_add_rPr()
            lang_el = rpr.find(qn("w:lang"))
            if lang_el is None:
                lang_el = OxmlElement("w:lang")
                rpr.append(lang_el)
            lang_el.set(qn("w:val"), lang)


def configure_styles(document):
    sec = document.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    if "Subtitle" not in styles:
        styles.add_style("Subtitle", WD_STYLE_TYPE.PARAGRAPH)

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(18, 52, 86)

    subtitle = styles["Subtitle"]
    subtitle.font.name = "Arial"
    subtitle.font.size = Pt(11)
    subtitle.font.italic = False
    subtitle.font.color.rgb = RGBColor(90, 102, 120)

    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(25, 55, 95)

    set_document_language(document)


def add_header_footer(document, label):
    for section in document.sections:
        header = section.header
        p = header.paragraphs[0]
        p.text = label
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.runs[0].font.name = "Arial"
        p.runs[0].font.size = Pt(9)
        p.runs[0].font.color.rgb = RGBColor(100, 116, 139)

        footer = section.footer
        p2 = footer.paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p2.add_run("Pagina ")
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(100, 116, 139)
        add_page_number(p2)


def add_cover(document, title, subtitle, meta_rows):
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(title)

    p = document.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(subtitle)

    document.add_paragraph("")
    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    for left, right in meta_rows:
        row = table.add_row().cells
        row[0].width = Inches(2.3)
        row[1].width = Inches(4.7)
        row[0].text = left
        row[1].text = right
        row[0].paragraphs[0].runs[0].bold = True
        for c in row:
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    document.add_paragraph("")


def add_heading(document, text, level=1):
    document.add_paragraph(text, style=f"Heading {level}")


def add_paragraph(document, text, bold_prefix=None):
    p = document.add_paragraph(style="Normal")
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        first.bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_capture_note(document, text):
    p = document.add_paragraph(style="Normal")
    run = p.add_run(f"[Sugerencia de captura: {text}]")
    run.italic = True
    run.bold = True
    run.font.color.rgb = RGBColor(128, 0, 0)
    return p


def add_bullets(document, items):
    for item in items:
        p = document.add_paragraph(style="Normal")
        p.style = document.styles["Normal"]
        p.paragraph_format.left_indent = Inches(0.35)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        p.add_run("• ").bold = True
        p.add_run(item)


def add_table(document, headers, rows, widths=None):
    table = document.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr_cells[idx].text = str(header)
        hdr_cells[idx].paragraphs[0].runs[0].bold = True
        hdr_cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(hdr_cells[idx], "DCE6F1")
        if widths and idx < len(widths):
            hdr_cells[idx].width = widths[idx]
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if widths and idx < len(widths):
                cells[idx].width = widths[idx]
    document.add_paragraph("")
    return table


def page_break(document):
    document.add_page_break()


def save_document(document, filename):
    path = OUTPUT_DIR / filename
    document.save(path)
    return path


def build_acta():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "Acta de Proyecto - CaniVet")
    add_cover(
        doc,
        "Acta de Proyecto",
        "Sistema de Gestion Canina y Veterinaria CaniVet",
        [
            ("Nombre del proyecto", PROJECT["nombre"]),
            ("Tipo de solucion", PROJECT["tipo"]),
            ("Tecnologia principal", f'{PROJECT["frontend"]} / {PROJECT["backend"]}'),
            ("Base de datos", PROJECT["bd"]),
            ("Autenticacion", PROJECT["auth"]),
            ("Fecha", "Mayo 2026"),
        ],
    )
    add_heading(doc, "1. Antecedentes", 1)
    add_paragraph(doc, "CaniVet surge como respuesta a la necesidad de digitalizar la operacion diaria de clinicas y centros de cuidado canino. En muchos entornos pequeños y medianos, la informacion de clientes, mascotas, citas, pagos, inventario y suscripciones se mantiene en hojas de calculo, cuadernos o conversaciones aisladas, lo que provoca duplicidad de datos, perdida de trazabilidad y tiempos de respuesta lentos.")
    add_paragraph(doc, "El proyecto se plantea como una plataforma web centralizada que integre la gestion administrativa y operativa del negocio veterinario en una sola interfaz, permitiendo trabajar con sucursales, roles de usuario, reportes y notificaciones por correo.")
    add_heading(doc, "2. Descripcion institucional y contexto", 1)
    add_paragraph(doc, "La propuesta se ubica dentro del contexto de una organizacion dedicada a servicios veterinarios, cuidado preventivo, control de citas, atencion al cliente, cobros, seguimiento de mascotas y servicios complementarios como guarderia y paseos. Estos procesos, aunque relacionados entre si, suelen administrarse de manera separada, con criterios distintos segun la persona o la sucursal que los ejecute.")
    add_paragraph(doc, "Desde una perspectiva academica y profesional, el proyecto busca demostrar la aplicacion de buenas practicas de desarrollo, orientacion a objetos, separacion por capas, seguridad de acceso, validacion de informacion y documentacion completa. Por ello, el sistema no se limita a una sola operacion, sino que integra varios procesos administrativos y de atencion en una plataforma coherente.")
    add_heading(doc, "3. Problema identificado", 1)
    add_paragraph(doc, "La gestion manual y fragmentada de la informacion genera inconsistencias entre registros, dificulta el seguimiento historico de mascotas y servicios, retrasa la facturacion y complica el control del inventario. Ademas, la ausencia de filtros y reportes oportunos limita la capacidad de tomar decisiones basadas en datos.")
    add_paragraph(doc, "Tambien se observa que la falta de controles estandarizados sobre autenticacion y permisos puede exponer informacion sensible o permitir errores operativos. Si no existe un sistema centralizado, una misma mascota puede quedar registrada con informacion incompleta, una cita puede no relacionarse correctamente con el cliente y los pagos pueden carecer de trazabilidad.")
    add_heading(doc, "4. Justificacion", 1)
    add_paragraph(doc, "La implementacion de CaniVet ofrece una solucion tecnologica capaz de organizar los procesos criticos del negocio veterinario. Con un sistema centralizado se reduce la dependencia del papel, se fortalece la seguridad de acceso, se mejora la calidad de la informacion y se agilizan tareas como programar citas, consultar historiales, registrar pagos, controlar stock y atender solicitudes de clientes.")
    add_paragraph(doc, "La justificacion del proyecto tambien se apoya en la necesidad de mejorar la experiencia de atencion. Una plataforma bien disenada permite responder con mayor rapidez, evitar la repeticion innecesaria de datos, mantener mejor organizadas las operaciones y obtener evidencia clara del comportamiento del negocio a traves de reportes.")
    add_heading(doc, "5. Objetivo general", 1)
    add_paragraph(doc, "Desarrollar una plataforma web orientada a la gestion integral de clinicas y centros de cuidado canino, capaz de administrar usuarios, clientes, mascotas, citas, pagos, inventario, suscripciones y reportes, con autenticacion segura, validaciones de datos y soporte de notificaciones por correo electronico.")
    add_heading(doc, "6. Objetivos especificos", 1)
    add_bullets(doc, [
        "Implementar autenticacion y control de acceso por roles mediante Supabase Auth.",
        "Desarrollar operaciones CRUD para clientes, mascotas, citas, servicios, pagos e inventario.",
        "Permitir la gestion de modulos complementarios como guarderia, paseos y suscripciones.",
        "Construir reportes exportables con filtros por sucursal, fecha y tipo de operacion.",
        "Incorporar validaciones de entrada y salida para proteger la integridad de la informacion.",
        "Automatizar correos para contacto, confirmaciones y comunicaciones operativas.",
        "Documentar el sistema desde la perspectiva funcional, tecnica y de analisis y diseno.",
    ])
    add_heading(doc, "7. Alcance", 1)
    add_paragraph(doc, "La primera version de CaniVet contempla modulos de autenticacion, panel administrativo, clientes, mascotas, citas, servicios, pagos, inventario, suscripciones, guarderia, paseos, auditoria, configuracion, reportes y formularios de contacto. La aplicacion web permite operar desde navegador, con frontend en React y backend Flask conectado a Supabase.")
    add_paragraph(doc, "No se incluye en esta fase una aplicacion movil nativa ni integracion completa con facturacion fiscal local. El componente de pagos externos se limita a integracion con Stripe Payment Links.")
    add_paragraph(doc, "El alcance operativo del sistema tambien incluye modulos de apoyo que fortalecen la propuesta integral del proyecto. Entre ellos se encuentran la administracion de sucursales, el directorio de usuarios con permisos diferenciados, el manejo de notificaciones internas, el registro de vacunas, el historial clinico por mascota, la generacion de facturas, el seguimiento de pagos online y el almacenamiento de fotos de servicio. Estos elementos permiten que el sistema se acerque a un escenario de uso real y no se limite a un CRUD academico basico.")
    add_heading(doc, "8. Limitaciones y restricciones", 1)
    add_paragraph(doc, "Como toda primera version, el proyecto presenta restricciones definidas para mantener la viabilidad del alcance. La solucion se enfoca en operacion web, por lo que no contempla una aplicacion movil nativa ni sincronizacion offline. Del mismo modo, aunque integra cobros y reportes, no incluye un esquema completo de contabilidad formal o integracion con entidades fiscales locales.")
    add_paragraph(doc, "Otra limitacion importante es la dependencia de servicios externos como Supabase, Gmail SMTP y Stripe para algunas funciones especificas. Si uno de estos proveedores cambia politicas de acceso o presenta indisponibilidad temporal, determinadas operaciones podrian verse afectadas hasta reconfigurar el sistema.")
    add_heading(doc, "9. Metodologia de trabajo", 1)
    add_paragraph(doc, "El desarrollo del proyecto se alinea con una metodologia iterativa, en la cual se construyen modulos funcionales por entregas parciales. Este enfoque permite validar primero la estructura principal del sistema, continuar con los procesos administrativos mas criticos y luego incorporar servicios complementarios, reportes y documentacion.")
    add_paragraph(doc, "La metodologia aplicada favorece la correccion temprana de errores y el refinamiento continuo del producto. Tambien permite que las pruebas y la documentacion no se dejen para el final, sino que evolucionen junto con el sistema.")
    add_paragraph(doc, "En terminos practicos, la metodologia de trabajo combina levantamiento, construccion, retroalimentacion y ajuste. Esto quiere decir que el proyecto no se desarrolla de manera totalmente lineal, sino que cada modulo puede revisarse y mejorarse a medida que otros componentes revelan nuevas necesidades. Esta dinamica resulta especialmente util en un sistema como CaniVet, donde una decision de autenticacion, datos o interfaz puede impactar varios procesos al mismo tiempo.")
    add_heading(doc, "10. Factibilidad del proyecto", 1)
    add_paragraph(doc, "La factibilidad tecnica del proyecto es alta debido a que se apoya en tecnologias ampliamente conocidas y bien documentadas, como React, Flask, Supabase y PostgreSQL. Estas herramientas ofrecen una curva de adopcion razonable y cuentan con suficiente material de soporte para desarrollar, probar y mantener la aplicacion de manera controlada.")
    add_paragraph(doc, "Desde el punto de vista operativo, la solucion es factible porque responde a procesos reales de una clinica o centro canino: atencion al cliente, gestion de mascotas, agenda, cobros, inventario y seguimiento de servicios. La informacion requerida por el sistema coincide con datos que la organizacion ya maneja, aunque actualmente los administre de forma dispersa.")
    add_paragraph(doc, "En cuanto a la factibilidad economica, el proyecto se beneficia del uso de herramientas de bajo costo relativo o con planes accesibles. Esto facilita que la propuesta pueda ejecutarse en un entorno academico y, al mismo tiempo, evolucionar hacia una implementacion de uso real con una inversion gradual.")
    add_heading(doc, "11. Recursos del proyecto", 1)
    add_table(
        doc,
        ["Tipo", "Detalle"],
        [
            ("Humanos", "Desarrollador o equipo academico, tutor o docente y usuario de validacion."),
            ("Software", "React, Vite, Flask, Supabase, PostgreSQL, GitHub, Visual Studio Code."),
            ("Infraestructura", "Equipo de desarrollo, navegador moderno, cuenta de Supabase y cuenta Gmail SMTP."),
            ("Documentacion", "README, manual de usuario, manual tecnico, analisis y diseno, acta y plan de actividades."),
        ],
        widths=[Inches(2), Inches(4.8)],
    )
    add_heading(doc, "12. Beneficiarios del proyecto", 1)
    add_paragraph(doc, "Los beneficiarios directos del proyecto son los usuarios internos que interactuan con la informacion del negocio: personal administrativo, personal operativo, encargados de sucursal y responsables de control. Todos ellos obtienen una mejora en la organizacion de la informacion y una reduccion de tareas manuales repetitivas.")
    add_paragraph(doc, "De forma indirecta, tambien se benefician los clientes finales, ya que una mejor organizacion interna repercute en una atencion mas agil, una agenda mejor controlada, mayor trazabilidad sobre el historial de las mascotas y una comunicacion mas clara frente a pagos, citas y servicios.")
    add_heading(doc, "13. Entregables del proyecto", 1)
    add_table(
        doc,
        ["Entregable", "Descripcion"],
        [
            ("Sistema web funcional", "Aplicacion con frontend, backend, autenticacion, modulos CRUD y reportes."),
            ("Base de datos", "Migraciones SQL y estructura persistente en Supabase."),
            ("Manual de usuario", "Documento para la operacion funcional del sistema."),
            ("Manual tecnico", "Documento para instalacion, estructura y mantenimiento."),
            ("Analisis y diseno", "Documento de requerimientos, arquitectura y modelado."),
            ("Cronograma y plan de actividades", "Instrumentos de planificacion formal del proyecto."),
            ("README", "Resumen tecnico del proyecto para repositorio y despliegue."),
        ],
        widths=[Inches(2.4), Inches(4.4)],
    )
    add_heading(doc, "14. Riesgos y mitigacion", 1)
    add_table(
        doc,
        ["Riesgo", "Impacto", "Mitigacion"],
        [
            ("Errores de configuracion en Supabase", "Interrumpe autenticacion y persistencia", "Mantener variables de entorno validadas y scripts de migracion versionados."),
            ("Dependencia del correo SMTP", "Fallo en notificaciones", "Configurar cuenta dedicada y manejar reintentos o mensajes alternativos."),
            ("Crecimiento de modulos", "Mayor complejidad de mantenimiento", "Aplicar separacion por capas, hooks reutilizables y modularidad en frontend/backend."),
            ("Datos incompletos en formularios", "Registros inconsistentes", "Usar validadores en backend y formularios controlados en frontend."),
        ],
        widths=[Inches(2.2), Inches(2.1), Inches(2.5)],
    )
    add_heading(doc, "15. Supuestos del proyecto", 1)
    add_bullets(doc, [
        "Que la organizacion dispone de los datos minimos necesarios para poblar los modulos principales.",
        "Que los usuarios finales cuentan con acceso basico a navegador y conectividad.",
        "Que el uso de Supabase, Flask y React se mantiene como stack principal durante el desarrollo.",
        "Que los procesos administrativos descritos no cambian radicalmente durante la fase de implementacion inicial.",
    ])
    add_heading(doc, "16. Criterios de exito", 1)
    add_bullets(doc, [
        "Que el sistema permita acceso autenticado y controlado por rol.",
        "Que los modulos principales operen correctamente con operaciones CRUD completas.",
        "Que existan reportes utilizables para consulta y exportacion.",
        "Que las validaciones eviten datos incompletos o inconsistentes.",
        "Que la documentacion final sea suficiente para uso, mantenimiento y defensa academica.",
    ])
    add_heading(doc, "17. Cronograma resumido", 1)
    add_table(
        doc,
        ["Semana", "Actividad principal", "Entregable"],
        [
            ("1", "Levantamiento y analisis de requerimientos", "Lista de necesidades y alcance"),
            ("2", "Diseno de base de datos y arquitectura", "Modelo de datos y estructura tecnica"),
            ("3", "Autenticacion y configuracion inicial", "Acceso, rutas y proteccion"),
            ("4-5", "CRUD y modulos operativos", "Clientes, mascotas, citas, servicios, pagos, inventario"),
            ("6", "Reportes, correo y ajustes", "Exportaciones, notificaciones y control de errores"),
            ("7-8", "Pruebas, documentacion y cierre", "Documentos finales y version entregable"),
        ],
        widths=[Inches(1), Inches(3.2), Inches(2.8)],
    )
    add_heading(doc, "18. Resultados esperados", 1)
    add_paragraph(doc, "Se espera disponer de una plataforma funcional que reduzca la dispersion de informacion, facilite la atencion de clientes, mejore el seguimiento de mascotas y permita visualizar indicadores operativos mediante reportes y paneles administrativos.")
    add_paragraph(doc, "A nivel academico, se espera que el proyecto evidencie el uso correcto de autenticacion, buenas practicas de desarrollo, enfoque orientado a objetos, documentacion tecnica y funcional, asi como integracion entre varias tecnologias.")
    add_paragraph(doc, "A nivel operativo, se espera que la plataforma permita administrar mejor escenarios multi-sucursal, mantener registros complementarios como vacunas e historial clinico, controlar servicios con mayor trazabilidad y ofrecer un mejor soporte para la toma de decisiones mediante paneles, exportaciones y auditoria.")
    add_heading(doc, "19. Conclusion", 1)
    add_paragraph(doc, "CaniVet representa una solucion integral y escalable para la gestion veterinaria. Su enfoque modular, el uso de autenticacion segura y la integracion entre frontend, backend y base de datos la convierten en una propuesta adecuada tanto para fines academicos como para evolucionar hacia un producto real de uso empresarial.")
    add_paragraph(doc, "Como acta de proyecto, este documento deja establecida una vision formal del trabajo a realizar, sus razones, sus limites, sus productos y el valor esperado. En ese sentido, funciona como punto de partida para la ejecucion, evaluacion y defensa del proyecto.")
    return save_document(doc, "Acta_Proyecto_CaniVet.docx")


def build_actividades():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "Plan de Actividades - CaniVet")
    add_cover(
        doc,
        "Plan de Actividades",
        "Cronograma de desarrollo y documentacion del proyecto CaniVet",
        [
            ("Proyecto", PROJECT["nombre"]),
            ("Metodologia", "Iterativa con entregas por modulo"),
            ("Duracion estimada", "8 semanas"),
            ("Fecha de elaboracion", "Mayo 2026"),
        ],
    )
    add_heading(doc, "1. Objetivo del plan", 1)
    add_paragraph(doc, "Organizar las actividades del proyecto CaniVet por fases, asignando resultados esperados y estableciendo una secuencia logica de trabajo desde el levantamiento de requerimientos hasta la entrega final de la plataforma y su documentacion.")
    add_paragraph(doc, "Este documento presenta una vision estructurada del trabajo necesario para ejecutar el proyecto. Su finalidad no es solo listar tareas, sino mostrar la secuencia metodologica, la dependencia entre actividades y la relacion entre cada fase de trabajo y los entregables finales.")
    add_heading(doc, "2. Fases del proyecto", 1)
    add_bullets(doc, [
        "Analisis de necesidades y definicion del alcance.",
        "Diseno de arquitectura, modelo de datos e interfaces.",
        "Desarrollo del frontend, backend y persistencia.",
        "Integracion de autenticacion, correo, reportes y pagos.",
        "Pruebas funcionales, correcciones y documentacion.",
    ])
    add_paragraph(doc, "La division por fases permite que el proyecto mantenga orden metodologico y que cada bloque de trabajo produzca resultados verificables antes de pasar al siguiente. Esto evita que la implementacion avance sin una base analitica suficiente y facilita identificar retrasos o dependencias no resueltas.")
    add_heading(doc, "3. Actividades por fase", 1)
    add_table(
        doc,
        ["No.", "Actividad", "Descripcion", "Responsable", "Duracion", "Entregable", "Estado"],
        [
            (1, "Levantamiento de informacion", "Identificacion de procesos clinicos, operativos y administrativos.", "Equipo de desarrollo", "3 dias", "Requerimientos iniciales", "Completado"),
            (2, "Analisis del sistema", "Definicion de modulos, actores, flujos y reglas de negocio.", "Equipo de desarrollo", "3 dias", "Analisis funcional", "Completado"),
            (3, "Diseno de base de datos", "Modelado de tablas para clientes, mascotas, citas, pagos, inventario y usuarios.", "Equipo de desarrollo", "4 dias", "Migraciones SQL", "Completado"),
            (4, "Diseno de interfaz", "Estructura de layouts publicos y administrativos con navegacion lateral.", "Equipo de desarrollo", "4 dias", "Pantallas base", "Completado"),
            (5, "Configuracion de autenticacion", "Implementacion de login, registro, verificacion de sesion y rutas protegidas.", "Equipo de desarrollo", "4 dias", "Modulo de acceso", "Completado"),
            (6, "CRUD de clientes y mascotas", "Registro, consulta, actualizacion y eliminacion de datos principales.", "Equipo de desarrollo", "5 dias", "Modulos operativos", "Completado"),
            (7, "CRUD de citas y servicios", "Programacion de agenda, servicios y flujo de atencion.", "Equipo de desarrollo", "4 dias", "Agenda funcional", "Completado"),
            (8, "Pagos e inventario", "Control de cobros, stock y movimientos relacionados.", "Equipo de desarrollo", "5 dias", "Operaciones financieras y stock", "Completado"),
            (9, "Suscripciones, guarderia y paseos", "Extensiones de servicios complementarios.", "Equipo de desarrollo", "4 dias", "Servicios avanzados", "En progreso"),
            (10, "Reportes y auditoria", "Paneles, tablas exportables y vistas analiticas.", "Equipo de desarrollo", "4 dias", "Modulo de reportes", "En progreso"),
            (11, "Correo y contacto", "Envio de correos desde formularios y eventos del sistema.", "Equipo de desarrollo", "2 dias", "Notificaciones", "Completado"),
            (12, "Pruebas y correcciones", "Validacion de formularios, rutas y respuestas del backend.", "Equipo de desarrollo", "4 dias", "Version estable", "En progreso"),
            (13, "Documentacion final", "Preparacion de acta, manuales, README y analisis/diseno.", "Equipo de desarrollo", "4 dias", "Paquete documental", "Completado"),
        ],
        widths=[Inches(0.45), Inches(1.35), Inches(2.25), Inches(1.1), Inches(0.8), Inches(1.5), Inches(0.9)],
    )
    add_heading(doc, "4. Cronograma resumido por semanas", 1)
    add_table(
        doc,
        ["Semana", "Actividades clave"],
        [
            ("Semana 1", "Levantamiento de informacion, analisis del problema y definicion del alcance."),
            ("Semana 2", "Diseno de base de datos, arquitectura y prototipo de interfaz."),
            ("Semana 3", "Autenticacion, registro, login y control de acceso."),
            ("Semana 4", "CRUD de clientes, mascotas, citas y servicios."),
            ("Semana 5", "Pagos, inventario, contacto y servicios complementarios."),
            ("Semana 6", "Reportes, exportaciones y auditoria."),
            ("Semana 7", "Pruebas funcionales, validaciones y ajustes."),
            ("Semana 8", "Documentacion, presentacion y entrega final."),
        ],
        widths=[Inches(1.2), Inches(5.6)],
    )
    add_heading(doc, "5. Criterios de control de actividades", 1)
    add_bullets(doc, [
        "Toda actividad debe producir un entregable o evidencia verificable.",
        "Las actividades de desarrollo deben apoyarse en validaciones y pruebas basicas.",
        "Las actividades de documentacion deben reflejar el estado real del proyecto.",
        "Las tareas criticas deben priorizarse sobre mejoras cosmeticas o extensiones no obligatorias.",
    ])
    add_paragraph(doc, "Estos criterios permiten convertir el plan de actividades en un instrumento de gestion y no solo en una lista descriptiva. La idea es que cada tarea pueda medirse, justificarse y relacionarse con un resultado concreto dentro del avance del proyecto.")
    add_heading(doc, "6. Dependencias entre actividades", 1)
    add_paragraph(doc, "Varias actividades del proyecto dependen directamente de otras. Por ejemplo, la construccion de reportes requiere que los modulos base ya produzcan datos consistentes; del mismo modo, la documentacion tecnica debe apoyarse en decisiones de arquitectura ya definidas y la documentacion de usuario necesita que las pantallas principales se encuentren razonablemente estables.")
    add_paragraph(doc, "Reconocer estas dependencias es importante porque evita una planificacion irreal. No todas las tareas pueden desarrollarse en paralelo con el mismo nivel de profundidad, y en proyectos administrativos suele ser preferible estabilizar primero el flujo principal antes de refinar extensiones.")
    add_heading(doc, "7. Observaciones finales", 1)
    add_paragraph(doc, "El plan de actividades refleja una implementacion incremental. La prioridad se concentra en los modulos obligatorios del proyecto: autenticacion, CRUD, reportes, filtrado de datos, envio de correos, buenas practicas y documentacion. Los servicios complementarios se integran sin comprometer la base administrativa del sistema.")
    add_paragraph(doc, "Las actividades de documentacion se consideran parte central del desarrollo y no una etapa secundaria. Por esa razon, la elaboracion del acta, manuales, README, analisis y cronograma se contempla como trabajo formal de cierre y sustentacion del proyecto.")
    add_paragraph(doc, "En resumen, este documento organiza el trabajo del proyecto de forma razonable y defendible. Sirve como apoyo para presentar el proceso de desarrollo, justificar tiempos y mostrar que la ejecucion del sistema responde a una secuencia de trabajo planificada.")
    return save_document(doc, "Plan_Actividades_CaniVet.docx")


def build_cronograma():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "Cronograma - CaniVet")
    add_cover(
        doc,
        "Cronograma del Proyecto",
        "Planificacion temporal formal para el desarrollo de CaniVet",
        [
            ("Proyecto", PROJECT["nombre"]),
            ("Horizonte", "8 semanas de trabajo"),
            ("Tipo de planificacion", "Cronograma por fases y entregables"),
            ("Fecha de emision", "Mayo 2026"),
        ],
    )
    add_heading(doc, "1. Introduccion", 1)
    add_paragraph(doc, "El cronograma del proyecto CaniVet establece la distribucion temporal de las tareas necesarias para completar el sistema de forma ordenada. Se ha organizado por semanas con el fin de mostrar la progresion del proyecto desde la definicion del problema hasta la documentacion final.")
    add_paragraph(doc, "La secuencia propuesta prioriza primero la comprension del problema y el diseno de la solucion, luego la implementacion de los modulos principales y, finalmente, la consolidacion mediante pruebas y documentacion. Esta estructura facilita el control del avance y la evaluacion de resultados en cada etapa.")
    add_heading(doc, "2. Criterios de planificacion", 1)
    add_bullets(doc, [
        "Las actividades se ordenan segun dependencia funcional y tecnica.",
        "Los modulos criticos se desarrollan antes que los complementarios.",
        "Las pruebas se ejecutan una vez existen flujos completos por modulo.",
        "La documentacion se formaliza al cierre, pero se alimenta durante todo el proceso.",
    ])
    add_heading(doc, "3. Cronograma detallado", 1)
    add_table(
        doc,
        ["Semana", "Periodo de trabajo", "Actividades", "Resultados esperados"],
        [
            ("1", "Inicio del proyecto", "Levantamiento de informacion, identificacion del problema, definicion del alcance y objetivos.", "Base conceptual clara del proyecto."),
            ("2", "Analisis y diseno", "Modelado de requerimientos, actores, casos de uso, estructura de base de datos y arquitectura general.", "Propuesta tecnica y funcional validada."),
            ("3", "Preparacion tecnica", "Configuracion del entorno, instalacion de dependencias, conexion con Supabase y estructura inicial del frontend/backend.", "Entorno listo para desarrollo."),
            ("4", "Seguridad y acceso", "Desarrollo del login, registro, gestion de sesion, roles y rutas protegidas.", "Modulo de autenticacion funcional."),
            ("5", "Modulos principales I", "CRUD de clientes, mascotas, citas y servicios.", "Flujo base de atencion operativo."),
            ("6", "Modulos principales II", "Pagos, inventario, suscripciones, guarderia, paseos y ajustes de integracion.", "Operacion administrativa ampliada."),
            ("7", "Analisis operativo", "Reportes, exportaciones, auditoria, filtros y correo electronico.", "Herramientas de consulta y seguimiento."),
            ("8", "Cierre del proyecto", "Pruebas funcionales, correccion de errores, elaboracion de documentacion y presentacion final.", "Version estable y paquete documental completo."),
        ],
        widths=[Inches(0.8), Inches(1.5), Inches(2.8), Inches(1.9)],
    )
    add_heading(doc, "4. Redaccion explicativa por etapa", 1)
    add_heading(doc, "4.1 Etapa de inicio", 2)
    add_paragraph(doc, "Durante las primeras semanas se concentra el esfuerzo de comprension del problema, delimitacion del alcance y analisis del entorno. Esta etapa es esencial porque define el marco de trabajo y evita que el sistema se construya sobre suposiciones poco claras.")
    add_heading(doc, "4.2 Etapa de construccion", 2)
    add_paragraph(doc, "Una vez definidos los requerimientos, se procede a la construccion del sistema por bloques. Se comienza con autenticacion y estructura general, luego se desarrollan los CRUD y finalmente se integran servicios de mayor valor analitico, como reportes, filtros y correo.")
    add_heading(doc, "4.3 Etapa de cierre", 2)
    add_paragraph(doc, "El cierre del cronograma no se limita a 'terminar de programar'. Incluye validar la calidad del producto, corregir inconsistencias y producir documentacion suficiente para sostener el proyecto desde el punto de vista tecnico, funcional y academico.")
    add_heading(doc, "5. Conclusion", 1)
    add_paragraph(doc, "El cronograma propuesto mantiene una relacion equilibrada entre analisis, desarrollo, pruebas y documentacion. Esto permite que CaniVet no sea solo una aplicacion operativa, sino tambien un proyecto formalmente sustentado y listo para presentacion.")
    return save_document(doc, "Cronograma_CaniVet.docx")


def build_manual_usuario():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "Manual de Usuario - CaniVet")
    add_cover(
        doc,
        "Manual de Usuario",
        "Guia funcional para operar la plataforma CaniVet",
        [
            ("Sistema", PROJECT["nombre"]),
            ("Tipo de usuario", "Administrador y usuario operativo"),
            ("Acceso", "Navegador web"),
            ("Version", "1.0"),
        ],
    )
    add_heading(doc, "1. Introduccion", 1)
    add_paragraph(doc, "Este manual explica el uso general de CaniVet desde la perspectiva del usuario final. La plataforma permite gestionar clientes, mascotas, citas, pagos, inventario, suscripciones y otros servicios asociados al cuidado canino desde una interfaz administrativa centralizada.")
    add_paragraph(doc, "Su objetivo es servir como guia practica para el personal que utilizara el sistema en sus actividades diarias. Por esa razon, cada modulo se describe en funcion de lo que el usuario necesita hacer dentro del sistema, las acciones que puede ejecutar y las recomendaciones para trabajar con mayor seguridad y orden.")
    add_heading(doc, "2. Objetivo del manual", 1)
    add_paragraph(doc, "Brindar instrucciones claras sobre el acceso, navegacion y uso de las funcionalidades principales de CaniVet, de modo que cualquier usuario autorizado pueda realizar sus tareas sin depender constantemente de asistencia tecnica.")
    add_heading(doc, "3. Requisitos de uso", 1)
    add_bullets(doc, [
        "Contar con conexion a internet o acceso al entorno local donde se publique el sistema.",
        "Utilizar un navegador actualizado como Google Chrome, Microsoft Edge o Mozilla Firefox.",
        "Disponer de credenciales validas de acceso registradas en el sistema.",
    ])
    add_heading(doc, "4. Acceso al sistema", 1)
    add_paragraph(doc, "El usuario debe ingresar a la direccion principal del sistema. Desde la pagina publica puede navegar por la pagina de inicio, el formulario de contacto y las pantallas de registro e inicio de sesion. Para autenticarse debe introducir correo electronico y contrasena. Si las credenciales son correctas, el sistema crea una sesion y redirige al area correspondiente.")
    add_paragraph(doc, "En un entorno real de trabajo, el acceso al sistema representa el primer control de seguridad. Por ello, cada usuario debe utilizar credenciales personales y evitar compartirlas. Una vez autenticado, el sistema adapta la experiencia segun los permisos y el contexto configurado.")
    add_capture_note(doc, "Insertar captura de la pagina de inicio de sesion o del formulario de acceso al sistema.")
    add_heading(doc, "5. Pantalla principal del administrador", 1)
    add_paragraph(doc, "Al entrar al panel administrativo, el usuario visualiza un dashboard con indicadores generales y un menu lateral con accesos organizados por secciones: principal, gestion, operaciones, servicios, analisis y sistema. La informacion visible depende del rol y de la sucursal activa.")
    add_capture_note(doc, "Insertar captura del dashboard principal o panel administrativo completo.")
    add_heading(doc, "6. Navegacion general", 1)
    add_paragraph(doc, "La navegacion se apoya en un panel lateral que agrupa las opciones de trabajo. Esta organizacion permite localizar rapidamente los modulos principales y mantener una experiencia consistente a lo largo del sistema. El usuario puede desplazarse entre dashboard, modulos de gestion, operaciones, servicios, analisis y configuracion.")
    add_paragraph(doc, "Ademas del menu lateral, cada pantalla suele presentar botones de accion, formularios, tablas y mensajes de confirmacion. Estos elementos se mantienen con un estilo uniforme para reducir la curva de aprendizaje.")
    add_capture_note(doc, "Insertar captura donde se aprecie el menu lateral, encabezado o navegacion general del sistema.")
    add_heading(doc, "7. Modulos del sistema", 1)
    add_paragraph(doc, "A continuacion se describen los modulos mas relevantes del sistema. La redaccion se centra en su uso practico, es decir, en como los utiliza el personal y para que sirven dentro de la operacion diaria.")
    add_heading(doc, "5.1 Clientes", 2)
    add_paragraph(doc, "En este modulo se registran los datos basicos de cada cliente, como nombre, correo, telefono y direccion. El usuario puede crear nuevos registros, editar informacion, eliminar elementos no vigentes y realizar busquedas por nombre o contacto.")
    add_paragraph(doc, "El modulo de clientes constituye uno de los puntos de entrada principales del sistema, ya que de estos registros dependen otros procesos como la vinculacion de mascotas, la programacion de citas, la facturacion y la comunicacion operativa. Por esa razon, se recomienda ingresar la informacion de manera completa y verificarla antes de guardar.")
    add_capture_note(doc, "Insertar captura del listado de clientes o del formulario para registrar un cliente.")
    add_heading(doc, "5.2 Mascotas", 2)
    add_paragraph(doc, "Permite vincular mascotas con sus propietarios, almacenando informacion como nombre, especie, raza, sexo, edad y peso. Desde aqui se consulta el historial relacionado y se mantiene actualizada la ficha de cada mascota.")
    add_paragraph(doc, "Ademas del registro basico, este modulo se relaciona con informacion complementaria como vacunas, historial clinico, citas, guarderia y paseos. Esto convierte a la ficha de mascota en un punto central para dar seguimiento integral al servicio prestado.")
    add_capture_note(doc, "Insertar captura del modulo de mascotas o de la ficha de una mascota registrada.")
    add_heading(doc, "5.3 Citas", 2)
    add_paragraph(doc, "El modulo de citas se utiliza para agendar atenciones, asignar fecha, hora, cliente, mascota y servicio. Tambien permite registrar notas y cambiar el estado de la cita segun avance la atencion.")
    add_paragraph(doc, "Una cita bien registrada permite mantener orden en la agenda de trabajo y evitar conflictos en la atencion. El usuario debe seleccionar correctamente al cliente, la mascota y el servicio, ya que estos datos impactan reportes, historial y operaciones posteriores.")
    add_capture_note(doc, "Insertar captura del calendario, listado de citas o formulario de programacion de citas.")
    add_heading(doc, "5.4 Servicios", 2)
    add_paragraph(doc, "El usuario administra el catalogo de servicios ofrecidos por la clinica o centro canino. Cada servicio puede incluir nombre, descripcion, precio y duracion estimada.")
    add_paragraph(doc, "La correcta configuracion de este modulo permite estandarizar la oferta de la empresa y facilita la asignacion rapida en el proceso de agenda y cobro. Tambien contribuye a que los reportes reflejen que servicios son los mas rentables o recurrentes.")
    add_capture_note(doc, "Insertar captura del modulo de servicios o del formulario de registro de servicios.")
    add_heading(doc, "5.5 Pagos", 2)
    add_paragraph(doc, "Desde pagos se registran cobros asociados a clientes y se lleva control del metodo y estado de pago. Esta informacion alimenta reportes financieros y exportaciones.")
    add_paragraph(doc, "En operaciones reales, este modulo ayuda a mantener trazabilidad sobre el flujo economico del negocio. Tambien sirve de apoyo para la emision de facturas y para el control de pagos online cuando estos se habilitan.")
    add_capture_note(doc, "Insertar captura del modulo de pagos, del registro de cobro o de la tabla de pagos.")
    add_heading(doc, "5.6 Inventario", 2)
    add_paragraph(doc, "En inventario se controlan productos, cantidades disponibles, precio y categoria. El sistema facilita identificar productos criticos o con bajo stock.")
    add_paragraph(doc, "Es recomendable revisar este modulo con frecuencia, ya que una mala actualizacion de stock puede distorsionar la lectura de reportes y afectar la disponibilidad de recursos para la atencion.")
    add_capture_note(doc, "Insertar captura del modulo de inventario o del listado de productos.")
    add_heading(doc, "5.7 Suscripciones, guarderia y paseos", 2)
    add_paragraph(doc, "Estos modulos amplian la oferta operativa del sistema y permiten gestionar servicios complementarios recurrentes y de cuidado especializado.")
    add_paragraph(doc, "Las suscripciones permiten organizar planes con fechas y proximos cobros; la guarderia registra entradas y salidas de mascotas en jornadas de cuidado; y los paseos documentan horario, duracion, distancia, responsable y observaciones. En conjunto, estos modulos muestran que la plataforma puede adaptarse a servicios mas alla de la consulta tradicional.")
    add_capture_note(doc, "Insertar captura de suscripciones, guarderia o paseos; puedes usar una o varias imagenes segun lo que quieras destacar.")
    add_heading(doc, "5.8 Vacunas e historial clinico", 2)
    add_paragraph(doc, "El sistema contempla almacenamiento de informacion clinica complementaria, como vacunas aplicadas, proximas dosis, motivo de consulta, sintomas, diagnostico, tratamiento y observaciones. Aunque esta informacion puede crecer con el tiempo, su estructura base ya permite una gestion mucho mas ordenada que un control manual disperso.")
    add_capture_note(doc, "Insertar captura del historial clinico o del registro de vacunas de una mascota.")
    add_heading(doc, "5.9 Notificaciones, auditoria y fotos de servicio", 2)
    add_paragraph(doc, "CaniVet incorpora funciones de apoyo a la operacion y supervison. Las notificaciones permiten registrar comunicaciones relevantes; la auditoria conserva evidencia de acciones sobre entidades; y las fotos de servicio sirven para documentar procesos antes y despues cuando el tipo de servicio lo requiere.")
    add_capture_note(doc, "Insertar captura de notificaciones, auditoria o fotos de servicio si esas pantallas estan implementadas y visibles.")
    add_heading(doc, "5.8 Reportes", 2)
    add_paragraph(doc, "La pantalla de reportes muestra indicadores, graficos y tablas exportables. Dependiendo del tipo de reporte, el usuario puede descargar datos en CSV, JSON o imprimir una version tipo PDF.")
    add_capture_note(doc, "Insertar captura del modulo de reportes, graficos o exportaciones disponibles.")
    add_heading(doc, "8. Dashboard", 1)
    add_paragraph(doc, "El dashboard muestra una vista resumida del estado general del sistema. Dependiendo del rol del usuario y de la configuracion de sucursal, puede presentar indicadores de ingresos, movimientos, notificaciones pendientes y accesos rapidos a los modulos mas usados. Su funcion es servir como punto de control inicial al comenzar la jornada.")
    add_paragraph(doc, "Desde la perspectiva del usuario, el dashboard no sustituye los modulos operativos, pero ayuda a interpretar rapidamente la situacion general. Es el punto adecuado para detectar si hay notificaciones, si existen pendientes o si algun modulo necesita ser revisado primero.")
    add_heading(doc, "9. Operaciones basicas", 1)
    add_bullets(doc, [
        "Registrar informacion: usar el boton de nuevo registro y completar los campos obligatorios.",
        "Editar informacion: seleccionar un elemento existente, modificar los datos y guardar.",
        "Eliminar informacion: confirmar la eliminacion solo cuando el registro ya no se requiera.",
        "Buscar y filtrar: utilizar los campos de busqueda y filtros por modulo, fecha o sucursal.",
        "Cerrar sesion: usar la opcion de salida ubicada en el panel lateral o superior.",
    ])
    add_paragraph(doc, "Estas operaciones, aunque parecen simples, son la base del funcionamiento diario. En una implementacion profesional, registrar, consultar, editar y filtrar datos con rapidez puede marcar la diferencia entre un sistema util y uno que solo se usa de manera parcial.")
    add_heading(doc, "10. Procedimientos frecuentes de uso", 1)
    add_paragraph(doc, "A continuacion se describen algunos procedimientos frecuentes que el usuario puede necesitar durante la operacion diaria del sistema.")
    add_bullets(doc, [
        "Registrar un cliente nuevo: ingresar al modulo de clientes, seleccionar la opcion de nuevo registro, completar los campos requeridos y guardar.",
        "Asociar una mascota a un cliente: abrir el modulo de mascotas, seleccionar el propietario correspondiente y completar la ficha de la mascota.",
        "Programar una cita: acceder al modulo de citas, indicar fecha, hora, cliente, mascota y servicio, luego confirmar el registro.",
        "Registrar un pago: dirigirse al modulo de pagos, seleccionar el cliente correspondiente, indicar monto, metodo y estado, y guardar.",
        "Consultar reportes: abrir el modulo de reportes, elegir el tipo de analisis requerido y aplicar los filtros necesarios.",
    ])
    add_capture_note(doc, "Insertar una captura por procedimiento importante si quieres explicar paso a paso como registrar o consultar informacion.")
    add_heading(doc, "11. Reportes y exportaciones", 1)
    add_paragraph(doc, "El sistema incluye un modulo de reportes que permite analizar informacion relevante como ingresos, metodos de pago, servicios rentables, comportamiento de clientes, datos por mascota, valor del inventario y comparativos por sucursal. El usuario puede revisar estos datos en pantalla y, segun el caso, exportarlos en formatos como CSV, JSON o impresiones tipo PDF.")
    add_paragraph(doc, "La utilidad de los reportes radica en que convierten informacion operativa en apoyo para la toma de decisiones. Por ejemplo, permiten conocer que servicios generan mayores ingresos, que clientes presentan mayor recurrencia y que productos del inventario requieren reposicion.")
    add_heading(doc, "12. Validaciones y mensajes", 1)
    add_paragraph(doc, "El sistema valida campos obligatorios, formatos de correo electronico, longitud maxima de texto y valores numericos minimos. Cuando un dato no cumple las reglas, el sistema muestra mensajes de error para evitar registros inconsistentes.")
    add_paragraph(doc, "El usuario debe prestar atencion a estos mensajes, ya que forman parte del mecanismo de calidad de datos. Guardar informacion incompleta o mal escrita puede afectar modulos relacionados, reportes y consultas futuras.")
    add_heading(doc, "13. Seguridad para el usuario", 1)
    add_paragraph(doc, "Cada usuario debe mantener la confidencialidad de sus credenciales y evitar compartir su sesion con terceros. El sistema aplica autenticacion y roles, pero la seguridad tambien depende de un uso responsable por parte del personal operativo.")
    add_heading(doc, "14. Recomendaciones de uso", 1)
    add_bullets(doc, [
        "Verificar la exactitud de los datos antes de guardar cambios.",
        "Mantener actualizado el inventario para evitar errores en reportes.",
        "Cerrar sesion al terminar la jornada, especialmente en equipos compartidos.",
        "Revisar periodicamente los reportes para detectar pendientes o anomalias.",
    ])
    add_paragraph(doc, "Se recomienda tambien mantener una disciplina operativa uniforme. Cuando varios usuarios trabajan en el mismo sistema, la consistencia en el registro de datos ayuda a que los resultados analiticos sean mas confiables y a que la informacion historica tenga verdadero valor de consulta.")
    add_heading(doc, "15. Preguntas frecuentes", 1)
    add_table(
        doc,
        ["Pregunta", "Respuesta"],
        [
            ("Que hago si no veo una opcion del menu?", "Lo mas probable es que tu rol no tenga permiso para esa vista o que la sucursal activa limite la informacion mostrada."),
            ("Puedo modificar un registro despues de guardarlo?", "Si tu perfil tiene permisos, puedes editarlo desde la tabla del modulo correspondiente."),
            ("Los reportes se actualizan automaticamente?", "En general, si. Los reportes dependen de la informacion registrada en el sistema y reflejan el estado actual de los datos disponibles."),
            ("Que pasa si cierro sesion?", "El sistema finaliza el acceso del usuario y sera necesario autenticarse nuevamente para volver al panel administrativo."),
        ],
        widths=[Inches(2.7), Inches(4.1)],
    )
    add_heading(doc, "16. Problemas frecuentes", 1)
    add_table(
        doc,
        ["Situacion", "Posible causa", "Accion recomendada"],
        [
            ("No puedo iniciar sesion", "Credenciales invalidas o sesion expirada", "Verificar correo y contrasena, luego intentar nuevamente."),
            ("No se guarda un formulario", "Campo obligatorio vacio o formato incorrecto", "Revisar los mensajes mostrados y corregir el dato."),
            ("No veo un modulo del menu", "Restriccion por rol o sucursal", "Consultar con el administrador del sistema."),
            ("No llega un correo", "Problema temporal con SMTP o direccion mal escrita", "Validar el correo del destinatario y revisar configuracion de envio."),
        ],
        widths=[Inches(1.8), Inches(2.1), Inches(2.5)],
    )
    add_heading(doc, "17. Glosario basico", 1)
    add_table(
        doc,
        ["Termino", "Significado dentro del sistema"],
        [
            ("CRUD", "Conjunto de operaciones para crear, leer, actualizar y eliminar registros."),
            ("Sucursal", "Ubicacion o sede activa desde la cual se contextualiza cierta informacion."),
            ("Dashboard", "Pantalla resumen con indicadores y accesos principales."),
            ("Reporte", "Salida de informacion consolidada para consulta o analisis."),
            ("Notificacion", "Aviso o mensaje interno relacionado con un evento del sistema."),
        ],
        widths=[Inches(1.8), Inches(5)],
    )
    add_heading(doc, "18. Conclusion", 1)
    add_paragraph(doc, "CaniVet ha sido disenado para ofrecer una experiencia de uso clara, organizada y profesional. Su correcta utilizacion contribuye a mejorar el servicio al cliente, el control de mascotas y la gestion administrativa general de la organizacion.")
    add_paragraph(doc, "Este manual debe considerarse una guia base. A medida que el sistema incorpore nuevos modulos o ajustes funcionales, sera recomendable actualizarlo para mantener alineada la operacion real con la documentacion entregada.")
    return save_document(doc, "Manual_Usuario_CaniVet.docx")


def build_manual_tecnico():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "Manual Tecnico - CaniVet")
    add_cover(
        doc,
        "Manual Tecnico",
        "Especificacion tecnica y de mantenimiento del sistema CaniVet",
        [
            ("Frontend", PROJECT["frontend"]),
            ("Backend", PROJECT["backend"]),
            ("Base de datos", PROJECT["bd"]),
            ("Autenticacion", PROJECT["auth"]),
            ("Version del documento", "1.0"),
        ],
    )
    add_heading(doc, "1. Descripcion general", 1)
    add_paragraph(doc, "CaniVet es una aplicacion web de arquitectura desacoplada. El frontend en React consume servicios HTTP del backend Flask y ambos dependen de Supabase para persistencia, autenticacion y consulta de datos. El sistema soporta operaciones administrativas de una clinica o centro de cuidado canino, con enfasis en modularidad y reutilizacion de componentes.")
    add_paragraph(doc, "Desde el punto de vista tecnico, la solucion fue concebida para separar claramente responsabilidades. La capa de presentacion gestiona la interaccion con el usuario; la capa de servicios concentra el acceso a API y proveedores externos; el backend valida y organiza la logica de negocio; y la capa de datos se apoya en Supabase para autenticacion y almacenamiento.")
    add_heading(doc, "2. Arquitectura", 1)
    add_table(
        doc,
        ["Capa", "Responsabilidad", "Tecnologia"],
        [
            ("Presentacion", "Interfaz publica y administrativa, navegacion, formularios y visualizacion de reportes.", "React, Vite, CSS"),
            ("Aplicacion", "Hooks, contexto, servicios HTTP y logica de interaccion con la API.", "React Context, hooks personalizados"),
            ("API", "Autenticacion, validacion, blueprints CRUD, contacto y pagos.", "Flask, flask-cors, requests"),
            ("Persistencia", "Autenticacion y base de datos en la nube.", "Supabase PostgreSQL"),
        ],
        widths=[Inches(1.2), Inches(3.9), Inches(1.4)],
    )
    add_capture_note(doc, "Insertar captura de un diagrama de arquitectura o esquema general del sistema si lo tienes.")
    add_heading(doc, "3. Estructura principal del proyecto", 1)
    add_bullets(doc, [
        "Canivet/src/components: layouts y componentes reutilizables de interfaz.",
        "Canivet/src/context: contexto de autenticacion y configuracion general.",
        "Canivet/src/hooks: hooks para clientes, mascotas, citas, inventario y utilidades CRUD.",
        "Canivet/src/pages: pantallas publicas y administrativas.",
        "Canivet/src/services: adaptadores de acceso a Supabase, backend y correo.",
        "backend/core: autenticacion, validadores y servicios base.",
        "backend/routes: blueprints para clientes, mascotas, citas, servicios, pagos, inventario, contacto, reservas, stripe y correo.",
        "migration.sql y versiones posteriores: scripts de base de datos en Supabase.",
    ])
    add_paragraph(doc, "Esta estructura responde a una intencion clara de mantenibilidad. En lugar de concentrar toda la logica en pocos archivos monoliticos, el proyecto distribuye responsabilidades para que cada parte del sistema tenga una funcion entendible y modificable. Eso facilita el trabajo de depuracion, pruebas y crecimiento funcional.")
    add_capture_note(doc, "Insertar captura del arbol de carpetas del proyecto desde el IDE o explorador de archivos.")
    add_heading(doc, "4. Seguridad y autenticacion", 1)
    add_paragraph(doc, "La autenticacion se soporta en Supabase Auth. El backend expone rutas /auth/login, /auth/register, /auth/me y /auth/redirect. El servicio de autenticacion obtiene el bearer token desde el encabezado Authorization o desde el cuerpo JSON, decodifica JWT firmados por HS256 o ES256 y resuelve el rol del usuario combinando metadatos del token con la lista de correos administrativos.")
    add_paragraph(doc, "Para proteger endpoints se emplean decoradores require_auth y require_admin. En frontend, la proteccion de rutas se implementa con ProtectedRoute y el contexto AuthContext, que mantiene la sesion sincronizada con Supabase.")
    add_heading(doc, "5. Validaciones", 1)
    add_paragraph(doc, "El backend incluye validadores por entidad en backend/core/validators.py. Se verifican campos requeridos, formato de correo electronico, valores numericos minimos, longitud maxima y existencia de contenido en actualizaciones parciales. Esto reduce errores y estandariza respuestas JSON cuando los datos no son validos.")
    add_table(
        doc,
        ["Entidad", "Campos validados", "Ejemplos de regla"],
        [
            ("Cliente", "nombre, email, telefono, direccion", "nombre requerido y email valido"),
            ("Mascota", "nombre, cliente_id, especie, raza, sexo, edad, peso", "peso numerico mayor o igual a 0"),
            ("Cita", "fecha, hora, cliente_id, mascota_id, servicio_id, estado", "todos los identificadores principales son obligatorios"),
            ("Servicio", "nombre, precio, descripcion, duracion", "precio requerido y no negativo"),
            ("Pago", "cliente_id, monto, metodo, estado", "monto requerido y no negativo"),
            ("Inventario", "nombre, cantidad, precio, descripcion, categoria", "cantidad entera no negativa"),
        ],
        widths=[Inches(1.2), Inches(2.7), Inches(2.4)],
    )
    add_heading(doc, "6. Integraciones", 1)
    add_bullets(doc, [
        "Supabase para autenticacion y almacenamiento de informacion.",
        "Gmail SMTP para envio de correos desde el backend.",
        "Stripe Payment Links para flujos externos de cobro.",
        "Chart.js para visualizacion de reportes en frontend.",
    ])
    add_paragraph(doc, "Estas integraciones no se agregan como accesorios aislados, sino como parte del flujo operativo. Supabase resuelve identidad y persistencia; Gmail SMTP apoya procesos de contacto; Stripe extiende los escenarios de cobro; y Chart.js transforma datos transaccionales en informacion visual util para gestion.")
    add_heading(doc, "7. Base de datos y modelo general", 1)
    add_paragraph(doc, "La base de datos se apoya en PostgreSQL administrado por Supabase. Las migraciones versionadas del proyecto permiten crear la estructura inicial, evolucionarla para reservas, configurar politicas y agregar control de usuarios del sistema. Este esquema facilita mantener trazabilidad sobre los cambios y reconstruir el entorno en otra instancia cuando sea necesario.")
    add_paragraph(doc, "Entre las entidades mas relevantes se encuentran clientes, mascotas, citas, servicios, pagos, inventario y usuarios_sistema. La relacion funcional entre ellas permite representar un flujo coherente: un cliente puede registrar una o varias mascotas; una mascota puede tener citas asociadas; cada cita se relaciona con un servicio; y los pagos permiten reflejar la operacion economica del negocio.")
    add_paragraph(doc, "Las migraciones tambien muestran la evolucion del proyecto hacia escenarios mas ricos, incorporando tablas para vacunas, historial clinico, suscripciones, guarderia, paseos, facturas, pagos online, notificaciones, auditoria, sucursales y fotos de servicio. Esto evidencia que la solucion fue pensada como una plataforma extensible.")
    add_table(
        doc,
        ["Tabla", "Proposito funcional"],
        [
            ("vacunas", "Registrar vacunas aplicadas, proximas dosis y datos del veterinario."),
            ("historial_clinico", "Conservar informacion de consultas, diagnosticos y tratamientos."),
            ("suscripciones", "Controlar planes recurrentes y proximos cobros."),
            ("guarderia", "Registrar entradas y salidas de mascotas en jornadas de cuidado."),
            ("paseos", "Gestionar horarios, rutas y responsables de paseos."),
            ("facturas", "Mantener evidencia estructurada de cobros emitidos."),
            ("pagos_online", "Dar seguimiento a enlaces de pago y estados remotos."),
            ("notificaciones", "Conservar mensajes y eventos relevantes del sistema."),
            ("auditoria", "Trazar acciones realizadas sobre entidades."),
            ("fotos_servicio", "Guardar evidencia visual asociada a una cita o servicio."),
        ],
        widths=[Inches(2.2), Inches(4.6)],
    )
    add_heading(doc, "8. Instalacion y ejecucion", 1)
    add_table(
        doc,
        ["Componente", "Pasos esenciales"],
        [
            ("Frontend", "Entrar a Canivet, ejecutar npm install y luego npm run dev."),
            ("Backend", "Entrar a backend, crear/activar entorno virtual, instalar requirements y ejecutar python app.py."),
            ("Variables de entorno", "Configurar credenciales de Supabase, JWT, correos SMTP y URL de API."),
            ("Base de datos", "Ejecutar migration.sql y sus versiones en Supabase SQL Editor."),
        ],
        widths=[Inches(1.5), Inches(5.3)],
    )
    add_paragraph(doc, "La instalacion del sistema requiere consistencia entre frontend, backend y base de datos. Si una de estas partes se configura con valores incorrectos, el comportamiento general puede degradarse incluso cuando el codigo se encuentra bien estructurado. Por esta razon, la documentacion de variables de entorno y migraciones tiene un papel central.")
    add_capture_note(doc, "Insertar captura de la consola ejecutando frontend o backend, o de las variables de entorno si deseas evidenciar la configuracion.")
    add_heading(doc, "9. Requisitos tecnicos", 1)
    add_table(
        doc,
        ["Categoria", "Detalle recomendado"],
        [
            ("Sistema operativo", "Windows 10 o superior, o entornos equivalentes para desarrollo web."),
            ("Navegador", "Chrome, Edge o Firefox en versiones modernas."),
            ("Node.js", "Version 18 o superior para ejecucion del frontend."),
            ("Python", "Version 3.12 o superior para backend Flask."),
            ("Base de datos", "Instancia activa de Supabase con migraciones aplicadas."),
            ("Herramientas", "VS Code o IDE equivalente, Git y acceso a consola."),
        ],
        widths=[Inches(2), Inches(4.8)],
    )
    add_heading(doc, "10. API y blueprints", 1)
    add_paragraph(doc, "El archivo backend/app.py actua como punto de entrada del servidor Flask. Ademas de las rutas de salud y autenticacion, registra blueprints especificos para clientes, mascotas, citas, servicios, pagos, inventario, contacto, reservas, stripe y correo. Esta distribucion mejora la mantenibilidad y evita concentrar toda la logica HTTP en un solo archivo.")
    add_paragraph(doc, "Cada blueprint atiende un dominio funcional particular y puede incorporar decoradores de autorizacion, validaciones y respuestas JSON normalizadas. Esta estrategia hace mas sencillo agregar nuevos modulos o modificar reglas sin comprometer el resto del sistema.")
    add_paragraph(doc, "Las pruebas incluidas en backend/tests/test_app.py validan al menos endpoints esenciales como health, autenticacion sin token y flujo de contacto. Aunque se trata de una cobertura inicial, la presencia de estas pruebas refleja una intencion de control de calidad desde backend.")
    add_table(
        doc,
        ["Ruta o grupo", "Proposito tecnico"],
        [
            ("/auth/*", "Gestionar autenticacion, identidad y redireccionamiento segun rol."),
            ("/api/clientes", "Gestion CRUD de clientes."),
            ("/api/mascotas", "Gestion CRUD y operaciones asociadas a mascotas."),
            ("/api/citas", "Agenda y control de citas."),
            ("/api/servicios", "Catalogo y administracion de servicios."),
            ("/api/pagos", "Registro de cobros y procesos relacionados."),
            ("/api/inventario", "Control de productos y stock."),
            ("/api/contacto y /api/email", "Canales de comunicacion y notificaciones."),
        ],
        widths=[Inches(2.2), Inches(4.6)],
    )
    add_capture_note(doc, "Insertar captura del archivo app.py o de las rutas registradas en backend si quieres reforzar la explicacion tecnica.")
    add_heading(doc, "11. Frontend, rutas y estado", 1)
    add_paragraph(doc, "En el frontend, el enrutador principal define vistas publicas como inicio, login y registro, asi como el segmento administrativo bajo /admin. La proteccion de acceso se gestiona mediante ProtectedRoute, mientras que AuthContext mantiene el estado del usuario autenticado y provee operaciones de login, registro, logout y lectura de token.")
    add_paragraph(doc, "Adicionalmente, existen hooks especializados como useAppointments, useClients, useInventory y usePets, que facilitan la reutilizacion de logica por dominio. Esta organizacion reduce duplicacion y favorece la consistencia del comportamiento entre pantallas.")
    add_paragraph(doc, "Un componente especialmente relevante es AppConfigContext, ya que centraliza la logica de sucursales, preferencias, notificaciones, auditoria, directorio de usuarios, historial de mascotas, fotos de servicio y filtros contextuales. Desde el punto de vista tecnico, esto resuelve un problema importante: mantener coherencia entre multiples modulos que dependen del contexto operativo activo.")
    add_capture_note(doc, "Insertar captura de AppRouter, AuthContext o AppConfigContext desde el editor de codigo si necesitas evidenciar la estructura tecnica.")
    add_heading(doc, "12. Migraciones y evolucion del modelo", 1)
    add_paragraph(doc, "El proyecto incluye varias migraciones SQL que reflejan una evolucion progresiva del modelo. La migracion base construye las estructuras iniciales; otras versiones agregan reservas, seguridad RLS y usuarios del sistema. Este enfoque favorece la trazabilidad, ya que cada cambio estructural puede reconstruirse y justificarse.")
    add_paragraph(doc, "Especialmente relevantes son las migraciones que incorporan tablas de apoyo como vacunas, historial_clinico, suscripciones, guarderia, paseos, facturas, pagos_online, notificaciones, auditoria, sucursales y fotos_servicio. Estas adiciones muestran que el sistema fue creciendo alrededor de requerimientos funcionales concretos.")
    add_heading(doc, "13. Mantenimiento", 1)
    add_paragraph(doc, "Para mantener el sistema se recomienda versionar cambios por modulo, documentar nuevas migraciones, validar compatibilidad entre el frontend y la API y ejecutar pruebas basicas sobre autenticacion, formularios, reportes y correo. Tambien es importante revisar que las variables de entorno continúen vigentes despues de cambios en proveedores externos.")
    add_paragraph(doc, "Las tareas de mantenimiento deben contemplar ademas la revision de integridad de datos, el seguimiento de nuevas tablas introducidas en migraciones posteriores y la verificacion de permisos asociados a roles y sucursales. En sistemas administrativos, muchos errores no se originan por codigo roto, sino por configuracion incorrecta o datos inconsistentes.")
    add_capture_note(doc, "Insertar captura de pruebas, consola de validacion o archivos de migracion si quieres apoyar esta seccion.")
    add_heading(doc, "14. Despliegue y consideraciones operativas", 1)
    add_paragraph(doc, "En un entorno de despliegue, el sistema requiere que el frontend conozca la URL del backend y que el backend posea las credenciales necesarias para comunicarse con Supabase y servicios SMTP. Tambien debe considerarse la configuracion de CORS para permitir solicitudes legitimas desde el dominio o puerto autorizado.")
    add_paragraph(doc, "Adicionalmente, si se habilitan funciones como pagos online o correo electronico, debe revisarse la vigencia de credenciales y politicas de los proveedores externos. Muchos problemas de operacion en sistemas web modernos no se originan en la logica del codigo, sino en permisos o configuraciones vencidas.")
    add_heading(doc, "15. Buenas practicas aplicadas", 1)
    add_bullets(doc, [
        "Separacion por modulos tanto en frontend como en backend.",
        "Uso de validadores centralizados para consistencia de datos.",
        "Proteccion de rutas y control de acceso por rol.",
        "Configuracion sensible mediante variables de entorno.",
        "Versionado de migraciones y estructura de despliegue clara.",
    ])
    add_heading(doc, "16. Problemas comunes y soluciones", 1)
    add_table(
        doc,
        ["Problema", "Diagnostico", "Solucion sugerida"],
        [
            ("Error CORS", "Origen no configurado en backend", "Agregar dominio en CORS_ORIGINS."),
            ("Token invalido", "JWT expirado o secreto incorrecto", "Verificar SUPABASE_JWT_SECRET y renovar sesion."),
            ("No llegan correos", "Credenciales SMTP incorrectas", "Revisar SMTP_HOST, SMTP_USER y SMTP_PASS."),
            ("Pantallas sin datos", "Migraciones o tabla faltante", "Ejecutar scripts SQL pendientes y revisar nombres de tabla."),
        ],
        widths=[Inches(1.8), Inches(2.2), Inches(2.4)],
    )
    add_heading(doc, "17. Conclusion", 1)
    add_paragraph(doc, "El sistema presenta una base tecnica coherente y extensible. La separacion entre interfaz, logica de aplicacion, API y servicios externos facilita el mantenimiento y permite incorporar nuevas funciones sin rehacer el nucleo existente.")
    add_paragraph(doc, "En consecuencia, el manual tecnico no solo describe como ejecutar la aplicacion, sino tambien como entender su estructura para corregir errores, incorporar mejoras y defender tecnicamente las decisiones de arquitectura tomadas en el desarrollo del proyecto.")
    return save_document(doc, "Manual_Tecnico_CaniVet.docx")


def build_analisis_diseno():
    doc = Document()
    configure_styles(doc)
    add_header_footer(doc, "Analisis y Diseno - CaniVet")
    add_cover(
        doc,
        "Documentacion de Analisis y Diseno",
        "Especificacion funcional y de arquitectura del sistema CaniVet",
        [
            ("Proyecto", PROJECT["nombre"]),
            ("Enfoque", "Analisis funcional y diseno orientado a objetos"),
            ("Dominio", "Gestion veterinaria y cuidado canino"),
            ("Fecha", "Mayo 2026"),
        ],
    )
    add_heading(doc, "1. Descripcion del problema", 1)
    add_paragraph(doc, "Los centros de atencion veterinaria y de cuidado canino necesitan manejar informacion heterogenea: clientes, mascotas, historial de servicios, citas, pagos, inventario y sucursales. Cuando estos procesos se gestionan de manera aislada, aparecen inconsistencias, tiempos muertos y poca trazabilidad.")
    add_paragraph(doc, "El problema no radica unicamente en la existencia de muchos datos, sino en la necesidad de relacionarlos correctamente. Para prestar un servicio de calidad es indispensable saber quien es el cliente, que mascotas tiene registradas, que historial acumula cada una, que servicios se han brindado, que pagos se realizaron y que recursos del negocio se utilizaron en el proceso.")
    add_heading(doc, "2. Justificacion del sistema", 1)
    add_paragraph(doc, "El sistema se justifica porque integra en una sola plataforma procesos que normalmente se encuentran dispersos. Esta integracion mejora la calidad de la informacion, disminuye errores de coordinacion y fortalece la trazabilidad operativa. Tambien incrementa la capacidad de supervision mediante reportes y paneles de control.")
    add_heading(doc, "3. Objetivo general del analisis y diseno", 1)
    add_paragraph(doc, "Definir de forma estructurada los requerimientos, actores, reglas de negocio, arquitectura y componentes de CaniVet, de modo que el proyecto posea una base formal para su implementacion, evaluacion y mantenimiento.")
    add_heading(doc, "4. Requerimientos funcionales", 1)
    add_bullets(doc, [
        "El sistema debe permitir registro, inicio y cierre de sesion.",
        "El sistema debe permitir CRUD de clientes, mascotas, citas, servicios, pagos e inventario.",
        "El sistema debe permitir gestionar suscripciones, guarderia y paseos.",
        "El sistema debe emitir reportes con exportacion de datos.",
        "El sistema debe ofrecer formulario de contacto y notificaciones por correo.",
        "El sistema debe restringir acceso segun rol y contexto de sucursal.",
    ])
    add_paragraph(doc, "Estos requerimientos funcionales no deben entenderse como una lista aislada, sino como capacidades interdependientes. Por ejemplo, la gestion de mascotas carece de valor si no se vincula con clientes, historial clinico y agenda. De igual modo, los reportes solo son utiles si los modulos operativos alimentan informacion consistente.")
    add_heading(doc, "5. Requerimientos no funcionales", 1)
    add_bullets(doc, [
        "Interfaz profesional, responsive y facil de usar.",
        "Seguridad basada en autenticacion y validacion de entradas.",
        "Mantenibilidad mediante modulos desacoplados y codigo reutilizable.",
        "Disponibilidad sobre infraestructura web y soporte para crecimiento por sucursales.",
    ])
    add_paragraph(doc, "Los requerimientos no funcionales son especialmente importantes en un sistema de este tipo porque la experiencia del usuario y la confiabilidad de la informacion son tan relevantes como la mera existencia de funciones. Una aplicacion con muchas pantallas, pero sin claridad visual o con mala validacion de datos, no resolveria el problema de fondo.")
    add_heading(doc, "6. Actores del sistema", 1)
    add_table(
        doc,
        ["Actor", "Descripcion", "Funciones principales"],
        [
            ("Administrador", "Usuario con acceso total al panel.", "Gestiona configuracion, usuarios, reportes y modulos operativos."),
            ("Usuario operativo", "Personal que atiende operaciones diarias.", "Registra clientes, mascotas, citas, pagos e inventario."),
            ("Cliente externo", "Usuario del sitio publico.", "Consulta informacion, contacta la clinica o inicia procesos de reserva."),
        ],
        widths=[Inches(1.4), Inches(2.2), Inches(2.8)],
    )
    add_heading(doc, "7. Casos de uso principales", 1)
    add_bullets(doc, [
        "Autenticarse en el sistema.",
        "Registrar un cliente y asociarle mascotas.",
        "Programar una cita con servicio asignado.",
        "Registrar un pago y actualizar su estado.",
        "Consultar reportes financieros, por cliente, mascota, inventario o sucursal.",
        "Enviar mensajes de contacto o notificaciones administrativas.",
    ])
    add_paragraph(doc, "A estos casos de uso se suman otros que amplian el alcance de la solucion, como registrar vacunas, almacenar historial clinico, gestionar suscripciones periodicas, coordinar guarderia y documentar paseos. Esto muestra que el sistema no se limita al ciclo minimo de cliente-mascota-cita, sino que busca representar una operacion mas completa.")
    add_heading(doc, "8. Reglas de negocio", 1)
    add_bullets(doc, [
        "Toda mascota debe estar asociada a un cliente.",
        "Toda cita requiere fecha, hora, cliente, mascota y servicio.",
        "Los montos de pago y precios de servicio no pueden ser negativos.",
        "La visibilidad de modulos depende del rol del usuario.",
        "Los correos administrativos pueden elevar permisos a rol administrador segun configuracion.",
        "Los registros deben filtrarse segun el contexto de sucursal cuando aplique.",
    ])
    add_paragraph(doc, "Las reglas de negocio buscan mantener coherencia entre el modelo logico y la operacion real. Sin este conjunto de restricciones, el sistema podria aceptar datos que despues no tendrian valor practico, como citas sin mascota, pagos incoherentes o usuarios con permisos que no corresponden a su funcion.")
    add_heading(doc, "9. Analisis del flujo de informacion", 1)
    add_paragraph(doc, "El flujo principal inicia cuando un usuario autenticado accede al panel administrativo y registra o consulta informacion. El sistema recibe la accion del usuario, valida los datos desde la interfaz y luego ejecuta una peticion al backend o directamente a Supabase segun el servicio involucrado. Finalmente, la respuesta se procesa para actualizar tablas, formularios, indicadores o reportes.")
    add_paragraph(doc, "Este flujo se repite en modulos como clientes, mascotas, citas y pagos, aunque con diferentes reglas y dependencias. La existencia de componentes reutilizables, hooks y servicios permite que el procesamiento siga patrones consistentes y sea mas facil de mantener.")
    add_paragraph(doc, "En modulos mas avanzados, como notificaciones, auditoria, fotos de servicio, vacunacion e historial clinico, el flujo de informacion incorpora ademas procesos de soporte que enriquecen la operacion principal. Esto significa que los datos ya no solo se registran para existir, sino para contextualizar mejor la atencion y mejorar la capacidad de supervision.")
    add_heading(doc, "10. Diseno orientado a objetos", 1)
    add_paragraph(doc, "Aunque el sistema combina React en frontend y Flask en backend, el enfoque de diseno mantiene entidades claras y responsabilidades delimitadas. Los componentes de interfaz representan vistas reutilizables; los hooks concentran logica de acceso; los servicios encapsulan comunicacion externa; y el backend segmenta autenticacion, validacion y rutas por dominio.")
    add_table(
        doc,
        ["Componente logico", "Responsabilidad"],
        [
            ("AuthContext / authService", "Gestionar sesion, token y estado del usuario autenticado."),
            ("Hooks CRUD", "Encapsular operaciones de consulta y mutacion por dominio."),
            ("Validadores de entidad", "Normalizar y verificar payloads antes de persistir."),
            ("Blueprints Flask", "Exponer API REST para cada modulo."),
            ("AppConfigContext", "Resolver permisos, sucursales, notificaciones y filtros de contexto."),
        ],
        widths=[Inches(2.5), Inches(4.3)],
    )
    add_heading(doc, "11. Diseno de base de datos", 1)
    add_paragraph(doc, "El proyecto utiliza Supabase PostgreSQL. Entre las tablas clave se encuentran clientes, mascotas, citas, servicios, pagos, inventario y usuarios_sistema. La tabla usuarios_sistema, incorporada en migration_v4.sql, extiende el control de acceso con rol, estado y sucursales permitidas.")
    add_paragraph(doc, "Las migraciones adicionales refuerzan el caracter evolutivo del proyecto. La inclusion de tablas como vacunas, historial_clinico, notificaciones, auditoria y fotos_servicio muestra que el sistema fue ampliado para responder a necesidades progresivamente mas reales, sin perder la coherencia general de la arquitectura.")
    add_table(
        doc,
        ["Entidad", "Funcion dentro del dominio"],
        [
            ("clientes", "Registrar propietarios o responsables de las mascotas."),
            ("mascotas", "Representar el sujeto principal de atencion veterinaria o de cuidado."),
            ("citas", "Planificar servicios y consultas por fecha y hora."),
            ("pagos", "Controlar transacciones economicas asociadas a servicios."),
            ("vacunas", "Registrar informacion preventiva y proximas dosis."),
            ("historial_clinico", "Mantener evidencia de atenciones, diagnosticos y tratamientos."),
            ("auditoria", "Conservar trazabilidad de acciones sobre el sistema."),
            ("notificaciones", "Apoyar comunicacion y seguimiento de eventos operativos."),
        ],
        widths=[Inches(2.1), Inches(4.7)],
    )
    add_capture_note(doc, "Insertar captura del modelo entidad-relacion, tablas en Supabase o diagrama de base de datos.")
    add_heading(doc, "12. Arquitectura de navegacion", 1)
    add_paragraph(doc, "La aplicacion publica incluye pagina de inicio, login, registro, servicios y contacto. El panel administrativo se agrupa bajo la ruta /admin y usa ProtectedRoute para impedir acceso no autorizado. El menu lateral organiza las funciones en dashboard, gestion, operaciones, servicios, analisis y sistema.")
    add_capture_note(doc, "Insertar captura del menu de navegacion o un esquema de rutas/pantallas.")
    add_heading(doc, "13. Diseno de interfaces", 1)
    add_paragraph(doc, "La interfaz fue concebida con una orientacion administrativa y profesional. El objetivo principal del diseno no es solo mostrar informacion, sino permitir que el usuario la interprete y manipule con facilidad. Por ello, se utilizan paneles, tablas, botones, filtros y modales organizados con jerarquia visual clara.")
    add_paragraph(doc, "Las pantallas publicas y privadas responden a necesidades distintas. Las publicas priorizan acceso, presentacion y contacto; las administrativas priorizan eficiencia, densidad informativa controlada y rapidez operativa.")
    add_capture_note(doc, "Insertar capturas de las pantallas principales del sistema, como login, dashboard, clientes, mascotas o reportes.")
    add_heading(doc, "14. Diseno de reportes y filtros", 1)
    add_paragraph(doc, "El modulo de reportes consolida informacion de clientes, citas, pagos, servicios e inventario. Admite filtros por sucursal y exportacion a CSV, JSON y formato imprimible. La visualizacion se apoya en Chart.js con tarjetas KPI, graficos de barras y dona, y tablas tabulares para analisis detallado.")
    add_paragraph(doc, "Desde la perspectiva de analisis y diseno, este modulo aporta valor porque traduce datos operativos en informacion sintetica util para gestion. No se trata solo de mostrar registros, sino de construir una capa de interpretacion que ayude a supervisar ingresos, preferencias de clientes, servicios dominantes y estado del inventario.")
    add_capture_note(doc, "Insertar captura del modulo de reportes y otra de los filtros aplicados si deseas mostrar el comportamiento analitico.")
    add_heading(doc, "15. Seguridad y autenticacion", 1)
    add_paragraph(doc, "La seguridad del sistema se sostiene sobre autenticacion, control de acceso, validacion de entradas y separacion de responsabilidades. La identidad del usuario se obtiene desde Supabase Auth y se complementa con logica interna para resolver permisos administrativos. Este enfoque evita exponer modulos sensibles a usuarios no autorizados.")
    add_paragraph(doc, "Las migraciones tambien evidencian una preocupacion por el control de acceso, como se observa en la configuracion de Row Level Security para tablas nuevas. Aunque las politicas actuales son permisivas en algunos casos, su existencia demuestra una base tecnica preparada para endurecer la seguridad en escenarios de produccion.")
    add_heading(doc, "16. Escalabilidad y mantenimiento", 1)
    add_paragraph(doc, "La separacion de concerns permite incorporar nuevas entidades o servicios. La existencia de hooks, servicios y blueprints por modulo facilita ampliar el sistema hacia auditoria avanzada, historiales clinicos mas ricos, automatizaciones por correo y nuevos esquemas de pago sin afectar de forma severa los modulos ya implementados.")
    add_paragraph(doc, "En terminos de diseno, esto significa que la solucion posee una estructura adecuada para crecer. La escalabilidad no depende solo de que la base de datos soporte mas registros, sino de que el codigo siga siendo legible, extensible y razonablemente desacoplado.")
    add_heading(doc, "17. Consideraciones de usabilidad", 1)
    add_paragraph(doc, "El sistema adopta una interfaz administrativa con panel lateral, jerarquia visual clara y organizacion por modulos. Esta decision no es meramente estetica; responde a la necesidad de que el usuario encuentre rapidamente las acciones principales, reduzca errores de navegacion y mantenga continuidad entre pantallas.")
    add_paragraph(doc, "La usabilidad tambien se apoya en formularios con validacion, mensajes de confirmacion y vistas tabulares para consulta. Estos elementos reducen la friccion durante la operacion y permiten que el sistema sea mas facil de aprender y sostener.")
    add_heading(doc, "18. Conclusion", 1)
    add_paragraph(doc, "Desde la perspectiva de analisis y diseno, CaniVet cumple con una estructura consistente para un proyecto academico-profesional. Integra procesos administrativos y de atencion en un ecosistema unificado, con foco en autenticacion, reportes, CRUD, usabilidad y mantenibilidad.")
    add_paragraph(doc, "El documento demuestra que el proyecto no fue construido de manera improvisada, sino a partir de una logica de analisis, modelado y diseno que puede defenderse tanto tecnica como funcionalmente.")
    return save_document(doc, "Analisis_Diseno_CaniVet.docx")


README_MD = """# CaniVet - Sistema de Gestion Canina y Veterinaria

## Descripcion

CaniVet es una plataforma web integral para la administracion de clinicas veterinarias y centros de cuidado canino. El sistema centraliza procesos de autenticacion, gestion de clientes, mascotas, citas, servicios, pagos, inventario, reportes y notificaciones por correo electronico.

El proyecto fue concebido como una solucion administrativa moderna que reduzca la dispersion de informacion y mejore la trazabilidad de cada operacion del negocio. Su arquitectura permite separar responsabilidades entre interfaz, API, autenticacion y persistencia, lo que facilita tanto el mantenimiento como la escalabilidad.

## Objetivo

Digitalizar y optimizar la operacion diaria del negocio veterinario, ofreciendo una solucion moderna, organizada y escalable que mejore la trazabilidad de la informacion y la calidad del servicio.

## Problema que resuelve

En muchas clinicas y centros de cuidado animal, la informacion se gestiona de forma manual o en herramientas dispersas. Esto genera errores de registro, dificultad para consultar historiales, lentitud para programar servicios y poca capacidad analitica para tomar decisiones. CaniVet responde a este problema mediante una plataforma centralizada que integra operaciones y datos relacionados.

## Caracteristicas principales

- Autenticacion con Supabase Auth.
- Panel administrativo con rutas protegidas.
- CRUD para clientes, mascotas, citas, servicios, pagos e inventario.
- Modulos adicionales de suscripciones, guarderia, paseos y auditoria.
- Reportes con exportacion a CSV, JSON y formato imprimible.
- Integracion de correo electronico mediante Gmail SMTP.
- Arquitectura desacoplada entre frontend React y backend Flask.

## Modulos del sistema

- Inicio, login y registro.
- Dashboard administrativo.
- Gestion de clientes.
- Gestion de mascotas.
- Agenda de citas.
- Catalogo de servicios.
- Gestion de pagos.
- Control de inventario.
- Suscripciones.
- Guarderia y paseos.
- Reportes y analitica.
- Auditoria y configuracion.

## Tecnologias utilizadas

- Frontend: React 19, Vite, CSS.
- Backend: Flask 3, flask-cors, requests.
- Base de datos: Supabase PostgreSQL.
- Autenticacion: Supabase Auth.
- Reportes: Chart.js.
- Pagos: Stripe Payment Links.
- Control de versiones: Git y GitHub.

## Estructura general

```text
canivet/
|-- Canivet/            # Frontend React
|-- backend/            # API Flask y logica de negocio
|-- docs/               # Documentacion y entregables
|-- migration.sql       # Scripts de base de datos
|-- migration_v2.sql
|-- migration_v3.sql
`-- migration_v4.sql
```

## Instalacion

### Frontend

```bash
cd Canivet
npm install
npm run dev
```

### Backend

```bash
cd backend
python -m venv .venv
.venv\\Scripts\\activate
pip install -r requirements.txt
python app.py
```

## Requisitos previos

- Node.js 18 o superior.
- Python 3.12 o superior.
- Cuenta y proyecto configurado en Supabase.
- Credenciales SMTP validas si se desea habilitar correo.
- Navegador moderno para acceso al panel.

## Variables de entorno importantes

- `VITE_SUPABASE_URL`
- `VITE_SUPABASE_ANON_KEY`
- `VITE_API_URL`
- `SUPABASE_URL`
- `SUPABASE_SERVICE_ROLE_KEY`
- `SUPABASE_JWT_SECRET`
- `SMTP_HOST`
- `SMTP_PORT`
- `SMTP_USER`
- `SMTP_PASS`

## Flujo general de uso

1. El usuario se registra o inicia sesion.
2. El sistema valida la identidad y determina el rol.
3. El usuario accede al panel administrativo.
4. Desde el panel puede gestionar clientes, mascotas, citas, servicios, pagos e inventario.
5. Los reportes permiten revisar indicadores y exportar resultados.
6. Las funciones de correo y contacto apoyan la comunicacion operativa.

## Buenas practicas aplicadas

- Separacion por capas y modulos.
- Uso de rutas protegidas.
- Validaciones de datos en backend.
- Variables de entorno para configuracion sensible.
- Versionado de migraciones y entregables.

## Documentacion incluida

- `Acta_Proyecto_CaniVet.docx`
- `Plan_Actividades_CaniVet.docx`
- `Cronograma_CaniVet.docx`
- `Manual_Usuario_CaniVet.docx`
- `Manual_Tecnico_CaniVet.docx`
- `Analisis_Diseno_CaniVet.docx`

## Estado del proyecto

Proyecto academico en desarrollo avanzado con modulos funcionales y paquete documental base listo para entrega.

## Mejoras futuras

- Historial clinico mas detallado.
- Mayor automatizacion de notificaciones.
- Version movil o adaptacion progresiva ampliada.
- Reportes gerenciales mas profundos.
- Integracion con procesos de facturacion formales.
"""


def write_readme():
    path = OUTPUT_DIR / "README_CaniVet.md"
    path.write_text(README_MD, encoding="utf-8")
    return path


def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    outputs = [
        build_acta(),
        build_actividades(),
        build_cronograma(),
        build_manual_usuario(),
        build_manual_tecnico(),
        build_analisis_diseno(),
        write_readme(),
    ]
    for item in outputs:
        print(item)


if __name__ == "__main__":
    main()
